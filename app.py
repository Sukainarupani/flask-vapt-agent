from flask import Flask, request, jsonify, send_file
import requests
import time
from docx import Document
from collections import defaultdict
import os

# ==============================
# ENVIRONMENT VARIABLES (CLOUD SAFE)
# ==============================
CLI_USERNAME = os.getenv("ZAP_USERNAME")
CLI_PASSWORD = os.getenv("ZAP_PASSWORD")
LOGIN_URL = os.getenv("LOGIN_URL")

# ==============================
# FLASK
# ==============================
app = Flask(__name__)
ZAP_API =  os.getenv("ZAP_API")
MAX_WAIT = 180

# ==============================
# HEALTH CHECK (RENDER NEEDS THIS)
# ==============================
@app.route("/submit", methods=["POST"])
def submit_from_n8n():
    data = request.json
    print("Received from n8n:", data)

    return jsonify({
        "status": "success",
        "received": data
    })
# ==============================
# HELPERS
# ==============================
def ensure_protocol(url):
    if not url.startswith("http://") and not url.startswith("https://"):
        return "https://" + url
    return url

def zap_request(endpoint, params=None):
    try:
        r = requests.get(f"{ZAP_API}{endpoint}", params=params, timeout=10)
        return r.json()
    except:
        print("[!] ZAP request failed:", endpoint)
        return None

# ==============================
# AUTH SETUP
# ==============================
def zap_setup_auth(target_url):
    if not CLI_USERNAME or not CLI_PASSWORD or not LOGIN_URL:
        print("[!] Auth skipped (missing environment variables)")
        return False

    print("[*] Setting up ZAP authentication...")

    context_name = "AuthContext"

    ctx_resp = zap_request("/JSON/context/action/newContext/", {
        "contextName": context_name
    })
    if not ctx_resp or "contextId" not in ctx_resp:
        print("[!] Failed to create context")
        return False

    context_id = ctx_resp["contextId"]

    zap_request("/JSON/context/action/includeInContext/", {
        "contextName": context_name,
        "regex": f"{target_url}.*"
    })

    auth_config = (
        f"loginUrl={LOGIN_URL}&"
        f"loginRequestData=username={CLI_USERNAME}&password={CLI_PASSWORD}"
    )

    zap_request("/JSON/authentication/action/setAuthenticationMethod/", {
        "contextId": context_id,
        "authMethodName": "formBasedAuthentication",
        "authMethodConfigParams": auth_config
    })

    user_resp = zap_request("/JSON/users/action/newUser/", {
        "contextId": context_id,
        "name": "cloud_user"
    })
    if not user_resp or "userId" not in user_resp:
        print("[!] Failed to create user")
        return False

    user_id = user_resp["userId"]

    zap_request("/JSON/users/action/setAuthenticationCredentials/", {
        "contextId": context_id,
        "userId": user_id,
        "authCredentialsConfigParams": f"username={CLI_USERNAME}&password={CLI_PASSWORD}"
    })

    zap_request("/JSON/users/action/setUserEnabled/", {
        "contextId": context_id,
        "userId": user_id,
        "enabled": "true"
    })

    zap_request("/JSON/forcedUser/action/setForcedUser/", {
        "contextId": context_id,
        "userId": user_id
    })

    zap_request("/JSON/forcedUser/action/setForcedUserModeEnabled/", {
        "boolean": "true"
    })

    print("[+] Authentication setup complete")
    return True

# ==============================
# SCANS
# ==============================
def zap_clear():
    zap_request("/JSON/core/action/deleteAllAlerts/")

def zap_spider(url):
    r = zap_request("/JSON/spider/action/scan/", {"url": url})
    scan_id = r.get("scan")
    start = time.time()
    while time.time() - start < MAX_WAIT:
        status = zap_request("/JSON/spider/view/status/", {"scanId": scan_id})
        if status and int(status.get("status", 0)) >= 100:
            break
        time.sleep(2)

def zap_active_scan(url):
    r = zap_request("/JSON/ascan/action/scan/", {"url": url})
    scan_id = r.get("scan")
    start = time.time()
    while time.time() - start < MAX_WAIT:
        status = zap_request("/JSON/ascan/view/status/", {"scanId": scan_id})
        if status and int(status.get("status", 0)) >= 100:
            break
        time.sleep(4)

def wait_for_passive():
    start = time.time()
    while time.time() - start < MAX_WAIT:
        status = zap_request("/JSON/pscan/view/recordsToScan/")
        if status and int(status.get("recordsToScan", 0)) == 0:
            break
        time.sleep(2)

def zap_alerts():
    r = zap_request("/JSON/core/view/alerts/")
    return r.get("alerts", []) if r else []

# ==============================
# 12-POINT MAPPING
# ==============================
def map_to_12_points(alerts, target_url):
    categories = {
        "Cleartext Transmission": [],
        "Session Timeout": ["session", "timeout"],
        "PHP Information Disclosure": ["php"],
        "Server Version Disclosure": ["server", "x-powered-by", "version"],
        "Private IP Disclosure": ["internal ip", "private ip", "localhost"],
        "Missing Security Headers": [
            "content security policy", "x-frame-options",
            "x-content-type-options", "strict-transport-security",
            "referrer-policy", "permissions-policy", "clickjacking"
        ],
        "Browser Cache Risk": ["cache-control", "pragma", "cache"],
        "Directory Listing": ["directory", "index of"],
        "Old jQuery": ["jquery"],
        "Missing CSRF Token": ["csrf"],
        "Cookie Hijacking Risk": ["cookie", "httponly", "secure", "samesite"],
        "Authorization Issues": ["authorization", "access control", "idor"]
    }

    results = {cat: {"status": "No Risk"} for cat in categories}

    if target_url.startswith("http://"):
        results["Cleartext Transmission"]["status"] = "Yes (Risk)"

    for alert in alerts:
        name = alert.get("alert", "").lower()
        for cat, keys in categories.items():
            for k in keys:
                if k in name:
                    results[cat]["status"] = "Yes (Risk)"

    return results

def build_12_point_findings(vapt_points):
    findings = []
    for name, info in vapt_points.items():
        detected = info["status"] != "No Risk"
        findings.append({
            "name": name,
            "status": "Detected" if detected else "Not Detected",
            "risk": "Medium" if detected else "No Risk",
            "description": f"{name} issue identified during scanning." if detected else f"No {name} related issue observed.",
            "observation": "Remediation is recommended." if detected else "No action required."
        })
    return findings

# ==============================
# REPORT GENERATION
# ==============================
def create_word_report(data, filename="VAPT_Report.docx"):
    doc = Document()
    doc.add_heading("VAPT Scan Report", level=1)

    doc.add_paragraph(f"Target URL: {data['target']}")
    doc.add_paragraph(f"Scan Type: {data['scan_type']}")
    doc.add_paragraph(f"Total Findings: {data['total_alerts']}")

    doc.add_heading("Detailed Findings", level=2)

    for f in data["raw_alerts"]:
        doc.add_heading(f["name"], level=3)
        doc.add_paragraph(f"Status: {f['status']}")
        doc.add_paragraph(f"Risk: {f['risk']}")
        doc.add_paragraph(f"Description: {f['description']}")
        doc.add_paragraph(f"Observation: {f['observation']}")

    doc.save(filename)
    return filename

# ==============================
# MAIN ROUTE
# ==============================
@app.route("/scan-download", methods=["POST"])
def scan_download():
    data = request.json
    url = data.get("url")

    if not url:
        return jsonify({"error": "URL missing"}), 400

    url = ensure_protocol(url)

    zap_clear()

    scan_type = "Unauthenticated"
    if zap_setup_auth(url):
        scan_type = "Authenticated"

    zap_spider(url)
    zap_active_scan(url)
    wait_for_passive()

    alerts = zap_alerts()
    vapt_points = map_to_12_points(alerts, url)
    findings = build_12_point_findings(vapt_points)

    report_data = {
        "target": url,
        "scan_type": scan_type,
        "total_alerts": sum(1 for f in findings if f["risk"] != "No Risk"),
        "raw_alerts": findings
    }

    filename = create_word_report(report_data)
    return send_file(filename, as_attachment=True)

# ==============================
# ENTRY POINT
# ==============================
if __name__ == "__main__":

    app.run(host="0.0.0.0", port=5000)
